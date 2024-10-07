# Ensure you have the required libraries
!pip install google-generativeai python-docx beautifulsoup4

import os
import google.generativeai as genai
from docx import Document
from docx.shared import RGBColor
from bs4 import BeautifulSoup  # To parse HTML

# Configure Gemini API
genai.configure(api_key="Your Api KEy")  # Replace YOUR_API_KEY with your actual API key

def generate_theory(aim):
    # Use Gemini-1.5-Flash to generate content for the Theory section
    model = genai.GenerativeModel("gemini-1.5-flash")

    # Input prompt for the theory section
    prompt = f"Generate a very large detailed theory section for the experiment with the aim: {aim}. Give me HTML code of this. The response should only involve HTML code, no explanation, nothing."

    # Generate content
    response = model.generate_content(prompt)

    # Return generated theory text in HTML format
    return response.text if response and response.text else "Error generating theory."

def html_to_docx(html_content, doc):
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process each element in the parsed HTML
    for element in soup.find_all(True):  # Find all tags
        if element.name == "h1":
            heading = doc.add_heading(element.text, level=1)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "h2":
            heading = doc.add_heading(element.text, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "h3":
            heading = doc.add_heading(element.text, level=3)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "h4":
            heading = doc.add_heading(element.text, level=4)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "h5":
            heading = doc.add_heading(element.text, level=5)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "h6":
            heading = doc.add_heading(element.text, level=6)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
        elif element.name == "p":
            doc.add_paragraph(element.text)
        elif element.name == "div":
            doc.add_paragraph(element.text)
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.text, style='ListNumber')
        elif element.name == "table":
            # Handle table conversion
            table = doc.add_table(rows=0, cols=0)
            for row in element.find_all("tr"):
                cells = row.find_all(["td", "th"])
                if len(cells) > len(table.columns):
                    # Adjust number of columns
                    for _ in range(len(cells) - len(table.columns)):
                        table.add_column()
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.get_text()

def create_experiment_doc(experiment_number, aim, theory, user_name, roll_number, batch):
    # Create a new Document
    doc = Document()

    # Add User Information at the top right
    user_info = doc.add_paragraph()
    user_info.add_run(f"Name: {user_name}    Roll No: {roll_number}    Batch: {batch}")
    user_info.alignment = 2  # Align to the right

    # Add Experiment Title (centered and bold)
    heading = doc.add_heading(f'Experiment - {experiment_number}', 0)  # Assign the heading object to 'heading'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Apply formatting to 'heading'

    # Add Aim section
    aim_heading = doc.add_heading('Aim', level=1)  # Assign the heading object to 'aim_heading'
    aim_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Apply formatting to 'aim_heading'
    doc.add_paragraph(aim)

    # Add Theory section
    html_to_docx(theory, doc)  # Convert and add HTML content to the doc

    # Save the document as 'experiment_aim_and_theory.docx'
    doc_name = f'experiment_{experiment_number}_aim_and_theory.docx'
    doc.save(doc_name)
    print(f"Document '{doc_name}' has been created successfully!")

if __name__ == "__main__":
    # Take Experiment Number, Aim, and User Info as input from the command line
    experiment_number = input("Enter the Experiment Number: ")
    aim = input("Enter the Aim: ")
    user_name = input("Enter your Name: ")
    roll_number = input("Enter your Roll Number: ")
    batch = input("Enter your Batch: ")

    # Generate Theory based on the Aim (using Gemini API)
    theory = generate_theory(aim)

    # Create the experiment document with Experiment Number, Aim, and Theory
    create_experiment_doc(experiment_number, aim, theory, user_name, roll_number, batch)
